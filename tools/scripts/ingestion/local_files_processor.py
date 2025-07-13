"""
Local Files Processor
Обработка локальных файлов из директории bootstrap
"""

import asyncio
import aiofiles
from pathlib import Path
from typing import Dict, List, Any, AsyncGenerator, Optional
from dataclasses import dataclass
import structlog
import hashlib
from datetime import datetime

# Опциональный импорт magic
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    magic = None

logger = structlog.get_logger()


@dataclass
class LocalDocument:
    """Локальный документ"""
    id: str
    title: str
    content: str
    file_path: str
    file_name: str
    file_extension: str
    file_size: int
    file_type: str
    created_date: str
    modified_date: str
    content_hash: str


class LocalFilesProcessor:
    """Процессор локальных файлов"""
    
    def __init__(self, config: Dict[str, Any], processing_config: Dict[str, Any]):
        self.config = config
        self.processing_config = processing_config
        
        # Настройки
        self.supported_formats = set(config.get("supported_formats", [".pdf", ".txt", ".epub", ".md", ".rst"]))
        self.max_file_size_mb = config.get("max_file_size_mb", 50)
        self.max_file_size_bytes = self.max_file_size_mb * 1024 * 1024
        self.encoding = config.get("encoding", "utf-8")
        
        # Инициализация libmagic для определения типов файлов
        self.mime = None
        if MAGIC_AVAILABLE:
            try:
                self.mime = magic.Magic(mime=True)
            except Exception as e:
                logger.warning("Failed to initialize libmagic", error=str(e))
                self.mime = None
        else:
            logger.warning("python-magic not available, file type detection will be limited")
    
    async def scan_directory(self, directory: Path) -> List[Path]:
        """Сканирование директории для поиска поддерживаемых файлов"""
        files = []
        
        try:
            for file_path in directory.rglob("*"):
                if file_path.is_file():
                    # Проверка расширения
                    file_extension = file_path.suffix.lower()
                    if file_extension in self.supported_formats:
                        # Проверка размера файла
                        try:
                            file_size = file_path.stat().st_size
                            if file_size <= self.max_file_size_bytes:
                                files.append(file_path)
                            else:
                                logger.warning(
                                    "File too large, skipping",
                                    file_path=str(file_path),
                                    size_mb=file_size / 1024 / 1024
                                )
                        except Exception as e:
                            logger.warning(
                                "Failed to get file stats",
                                file_path=str(file_path),
                                error=str(e)
                            )
            
            logger.info("Scanned directory", directory=str(directory), files_found=len(files))
            return files
            
        except Exception as e:
            logger.error("Failed to scan directory", directory=str(directory), error=str(e))
            return []
    
    async def process_file(self, file_path: Path) -> Optional[LocalDocument]:
        """Обработка отдельного файла"""
        try:
            # Получение информации о файле
            stat = file_path.stat()
            file_size = stat.st_size
            created_date = datetime.fromtimestamp(stat.st_ctime).isoformat()
            modified_date = datetime.fromtimestamp(stat.st_mtime).isoformat()
            
            # Определение типа файла
            file_type = "text/plain"
            if self.mime:
                try:
                    file_type = self.mime.from_file(str(file_path))
                except Exception:
                    pass
            
            # Извлечение содержимого в зависимости от типа файла
            content = await self._extract_content(file_path, file_type)
            
            if not content or len(content.strip()) < 50:
                logger.warning("File content too short or empty", file_path=str(file_path))
                return None
            
            # Создание хеша содержимого
            content_hash = hashlib.sha256(content.encode()).hexdigest()
            
            # Создание уникального ID
            file_id = f"local:{content_hash[:16]}"
            
            return LocalDocument(
                id=file_id,
                title=self._get_file_title(file_path),
                content=content,
                file_path=str(file_path),
                file_name=file_path.name,
                file_extension=file_path.suffix.lower(),
                file_size=file_size,
                file_type=file_type,
                created_date=created_date,
                modified_date=modified_date,
                content_hash=content_hash
            )
            
        except Exception as e:
            logger.error("Failed to process file", file_path=str(file_path), error=str(e))
            return None
    
    async def _extract_content(self, file_path: Path, file_type: str) -> str:
        """Извлечение содержимого файла в зависимости от типа"""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension in [".txt", ".md", ".rst"]:
                return await self._extract_text_content(file_path)
            elif file_extension == ".pdf":
                return await self._extract_pdf_content(file_path)
            elif file_extension == ".epub":
                return await self._extract_epub_content(file_path)
            elif file_extension == ".docx":
                return await self._extract_docx_content(file_path)
            else:
                # Попытка извлечь как текст
                return await self._extract_text_content(file_path)
                
        except Exception as e:
            logger.error(
                "Failed to extract content",
                file_path=str(file_path),
                file_type=file_type,
                error=str(e)
            )
            return ""
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Извлечение текстового содержимого"""
        try:
            async with aiofiles.open(file_path, 'r', encoding=self.encoding) as f:
                content = await f.read()
                return content.strip()
        except UnicodeDecodeError:
            # Попытка с другими кодировками
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        content = await f.read()
                        return content.strip()
                except UnicodeDecodeError:
                    continue
            raise
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Извлечение содержимого PDF файла"""
        try:
            import PyPDF2
            
            content = []
            
            async with aiofiles.open(file_path, 'rb') as f:
                pdf_data = await f.read()
                
            # Выполнение в отдельном потоке для избежания блокировки
            def extract_pdf_text():
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
                text_content = []
                
                for page in pdf_reader.pages:
                    try:
                        text_content.append(page.extract_text())
                    except Exception as e:
                        logger.warning("Failed to extract page text", error=str(e))
                
                return "\n".join(text_content)
            
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, extract_pdf_text)
            
            return content.strip()
            
        except Exception as e:
            logger.error("Failed to extract PDF content", file_path=str(file_path), error=str(e))
            return ""
    
    async def _extract_epub_content(self, file_path: Path) -> str:
        """Извлечение содержимого EPUB файла"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            def extract_epub_text():
                book = epub.read_epub(str(file_path))
                content = []
                
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        text = soup.get_text()
                        if text.strip():
                            content.append(text.strip())
                
                return "\n\n".join(content)
            
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, extract_epub_text)
            
            return content.strip()
            
        except Exception as e:
            logger.error("Failed to extract EPUB content", file_path=str(file_path), error=str(e))
            return ""
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Извлечение содержимого DOCX файла"""
        try:
            from docx import Document
            
            def extract_docx_text():
                doc = Document(str(file_path))
                content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                
                return "\n".join(content)
            
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, extract_docx_text)
            
            return content.strip()
            
        except Exception as e:
            logger.error("Failed to extract DOCX content", file_path=str(file_path), error=str(e))
            return ""
    
    def _get_file_title(self, file_path: Path) -> str:
        """Получение заголовка файла"""
        # Используем имя файла без расширения как заголовок
        title = file_path.stem
        
        # Очистка и форматирование заголовка
        title = title.replace("_", " ").replace("-", " ")
        title = " ".join(word.capitalize() for word in title.split())
        
        return title
    
    async def process_files_batches(self, directory: Path) -> AsyncGenerator[List[LocalDocument], None]:
        """Обработка файлов батчами"""
        try:
            # Сканирование директории
            files = await self.scan_directory(directory)
            
            if not files:
                logger.warning("No supported files found", directory=str(directory))
                return
            
            batch_size = self.processing_config.get("batch_size", 50)
            max_workers = self.processing_config.get("max_workers", 10)
            
            # Обработка файлов батчами
            for i in range(0, len(files), batch_size):
                batch_files = files[i:i + batch_size]
                
                # Ограничение количества параллельных задач
                semaphore = asyncio.Semaphore(max_workers)
                
                async def process_with_semaphore(file_path):
                    async with semaphore:
                        return await self.process_file(file_path)
                
                # Параллельная обработка файлов в батче
                tasks = [process_with_semaphore(file_path) for file_path in batch_files]
                results = await asyncio.gather(*tasks, return_exceptions=True)
                
                # Фильтрация успешно обработанных файлов
                valid_documents = []
                for result in results:
                    if isinstance(result, LocalDocument):
                        valid_documents.append(result)
                    elif isinstance(result, Exception):
                        logger.error("File processing failed", error=str(result))
                
                if valid_documents:
                    logger.info(
                        "Processed file batch",
                        batch_size=len(batch_files),
                        successful=len(valid_documents)
                    )
                    yield valid_documents
                
                # Небольшая пауза между батчами
                await asyncio.sleep(0.1)
                
        except Exception as e:
            logger.error("Failed to process files", directory=str(directory), error=str(e))
            raise 